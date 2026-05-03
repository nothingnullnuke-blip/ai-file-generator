from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

from generators.base_generator import BaseGenerator
from models.schemas import DataType
from utils.logger import setup_logger

logger = setup_logger(__name__)

class DocxGenerator(BaseGenerator):
    def __init__(self, structured_output, output_filename=None):
        super().__init__(structured_output, output_filename)
        self.doc = Document()
        self._set_document_style()
    
    def generate(self) -> Path:
        try:
            logger.info(f"Generating DOCX: {self.filepath}")
            
            self._add_title()
            
            if self.output.description:
                self._add_description()
            
            for section in sorted(self.output.sections, key=lambda x: x.order):
                self._add_section(section)
            
            self.doc.save(str(self.filepath))
            logger.info(f"DOCX generated: {self.filepath}")
            return self.filepath
            
        except Exception as e:
            logger.error(f"DOCX generation failed: {str(e)}")
            raise
    
    def get_file_extension(self) -> str:
        return "docx"
    
    def _set_document_style(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def _add_title(self):
        title = self.doc.add_paragraph(self.output.title, style='Heading 1')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_format = title.runs[0]
        title_format.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors['primary']))
        title_format.font.size = Pt(24)
        title_format.font.bold = True
    
    def _add_description(self):
        para = self.doc.add_paragraph(self.output.description)
        para.style = 'Normal'
        para.paragraph_format.space_after = Pt(12)
    
    def _add_section(self, section):
        if section.type == DataType.HEADING:
            self._add_heading(section)
        elif section.type == DataType.TEXT:
            self._add_text(section)
        elif section.type == DataType.TABLE:
            self._add_table(section)
    
    def _add_heading(self, section):
        text = section.content.get('text', '')
        heading = self.doc.add_paragraph(text, style='Heading 2')
        heading_format = heading.runs[0]
        heading_format.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors['secondary']))
        heading_format.font.size = Pt(16)
        heading_format.font.bold = True
    
    def _add_text(self, section):
        text = section.content.get('text', '')
        para = self.doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)
    
    def _add_table(self, section):
        content = section.content
        title = content.get('title', '')
        headers = content.get('headers', [])
        rows = content.get('rows', [])
        
        if title:
            self.doc.add_paragraph(title, style='Heading 3')
        
        table = self.doc.add_table(rows=len(rows)+1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            header_cells[idx].text = str(header)
        
        for row_idx, row_data in enumerate(rows, 1):
            row_cells = table.rows[row_idx].cells
            for col_idx, value in enumerate(row_data):
                row_cells[col_idx].text = str(value)
    
    @staticmethod
    def _hex_to_rgb(hex_color):
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))